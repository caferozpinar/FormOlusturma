"""
Şablon İşlemleri Modülü
=======================

Word şablonlarında placeholder değiştirme ve belge birleştirme.

YENİ: Geçici dosyaları otomatik temizleme özelliği eklendi.
"""

from __future__ import annotations

import logging
import re
import uuid
import atexit
import shutil
from pathlib import Path
from typing import Any, Iterable, Optional
from contextlib import contextmanager

from docx import Document as DocxDocument
from docxcompose.composer import Composer

# Modül günlükleyicisi
gunluk = logging.getLogger(__name__)

# Placeholder regex deseni
_PLACEHOLDER_RE = re.compile(r"\{/\s*([A-Za-z0-9_]+)\s*/\}")

# Geçici dosyaları takip için
_gecici_dosyalar: set[Path] = set()


def _gecici_dosya_kaydet(yol: Path) -> None:
    """Geçici dosyayı takip listesine ekler."""
    _gecici_dosyalar.add(yol)


def _tum_gecici_dosyalari_temizle() -> None:
    """Program sonunda tüm geçici dosyaları temizler."""
    for yol in list(_gecici_dosyalar):
        try:
            if yol.exists():
                yol.unlink()
                gunluk.debug(f"Geçici dosya silindi: {yol.name}")
        except Exception as e:
            gunluk.warning(f"Geçici dosya silinemedi {yol}: {e}")
    _gecici_dosyalar.clear()


# Program sonunda otomatik temizlik
atexit.register(_tum_gecici_dosyalari_temizle)


@contextmanager
def gecici_dosya_konteksi(gecici_dosya_yolu: Path):
    """
    Geçici dosya için context manager.

    Kullanım:
    --------
    with gecici_dosya_konteksi(yol) as dosya:
        # dosya ile işlem yap
        pass
    # Buraya gelince otomatik silinir
    """
    try:
        yield gecici_dosya_yolu
    finally:
        try:
            if gecici_dosya_yolu.exists():
                gecici_dosya_yolu.unlink()
                gunluk.debug(f"Geçici dosya temizlendi: {gecici_dosya_yolu.name}")
                _gecici_dosyalar.discard(gecici_dosya_yolu)
        except Exception as e:
            gunluk.warning(f"Geçici dosya temizlenemedi {gecici_dosya_yolu}: {e}")


def gecici_dizini_temizle(gecici_dizin: Path, gunluk_ref: Optional[logging.Logger] = None) -> None:
    """
    Geçici dizindeki tüm dosyaları temizler.

    Parametreler:
    -------------
    gecici_dizin : Path
        Temizlenecek dizin
    gunluk_ref : Logger, optional
        Logger referansı
    """
    if not gecici_dizin.exists():
        return

    try:
        dosya_sayisi = 0
        for dosya in gecici_dizin.glob("*"):
            if dosya.is_file():
                dosya.unlink()
                dosya_sayisi += 1

        if gunluk_ref and dosya_sayisi > 0:
            gunluk_ref.info(f"Geçici dizin temizlendi: {dosya_sayisi} dosya silindi")

    except Exception as e:
        if gunluk_ref:
            gunluk_ref.warning(f"Geçici dizin temizlenemedi: {e}")


def _tum_paragraflari_yinele(kapsayici: Any) -> Iterable:
    """Belge, tablo, header, footer içindeki tüm paragrafları yineler."""
    if hasattr(kapsayici, "paragraphs"):
        for p in kapsayici.paragraphs:
            yield p

    if hasattr(kapsayici, "tables"):
        for tablo in kapsayici.tables:
            for satir in tablo.rows:
                for hucre in satir.cells:
                    yield from _tum_paragraflari_yinele(hucre)


def _paragrafta_degistir(paragraf: Any, esleme: dict[str, str]) -> None:
    """Tek bir paragrafta placeholder'ları değiştirir (run-safe)."""
    runlar = paragraf.runs
    if not runlar:
        return

    # Tüm run metinlerini birleştir
    tam_metin = "".join(r.text or "" for r in runlar)
    if not tam_metin:
        return

    orijinal = tam_metin

    # Direkt anahtar değiştirme
    for anahtar, deger in esleme.items():
        if anahtar in tam_metin:
            tam_metin = tam_metin.replace(anahtar, str(deger))

    # Normalize edilmiş placeholder değiştirme ({/ KEY /} varyasyonları)
    def normalize_placeholder(esleme_obj):
        anahtar = esleme_obj.group(1)
        tam_anahtar = "{/" + anahtar + "/}"
        if tam_anahtar in esleme:
            return str(esleme[tam_anahtar])
        return esleme_obj.group(0)

    tam_metin = _PLACEHOLDER_RE.sub(normalize_placeholder, tam_metin)

    # Sadece değişiklik varsa güncelle
    if tam_metin != orijinal:
        runlar[0].text = tam_metin
        for r in runlar[1:]:
            r.text = ""


def belgede_placeholder_degistir(
    belge: DocxDocument,
    esleme: dict[str, str],
    gunluk_ref: Optional[logging.Logger] = None
) -> None:
    """
    Belge gövdesi, tablolar, header ve footer'lardaki tüm placeholder'ları değiştirir.

    Parametreler:
    -------------
    belge : DocxDocument
        İşlenecek belge
    esleme : dict[str, str]
        {/PLACEHOLDER/} -> değer eşlemesi
    gunluk_ref : Logger, optional
        Loglama için logger referansı
    """
    degistirilen = 0

    # Gövde ve tablolar
    for paragraf in _tum_paragraflari_yinele(belge):
        onceki = "".join(r.text or "" for r in paragraf.runs)
        _paragrafta_degistir(paragraf, esleme)
        sonraki = "".join(r.text or "" for r in paragraf.runs)
        if onceki != sonraki:
            degistirilen += 1

    # Header ve footer'lar
    for bolum in belge.sections:
        hf_listesi = (
            bolum.header, bolum.footer,
            bolum.first_page_header, bolum.first_page_footer,
            bolum.even_page_header, bolum.even_page_footer
        )
        for hf in hf_listesi:
            for paragraf in _tum_paragraflari_yinele(hf):
                onceki = "".join(r.text or "" for r in paragraf.runs)
                _paragrafta_degistir(paragraf, esleme)
                sonraki = "".join(r.text or "" for r in paragraf.runs)
                if onceki != sonraki:
                    degistirilen += 1

    if gunluk_ref:
        gunluk_ref.info(f"{degistirilen} paragrafta placeholder değiştirildi")


def gecici_dosyaya_render_et(
    kaynak_docx: Path,
    esleme: dict[str, str],
    gecici_dizin: Path,
    gunluk_ref: Optional[logging.Logger] = None,
    otomatik_takip: bool = True
) -> Path:
    """
    Şablonu placeholder değiştirme ile geçici dosyaya render eder.

    Parametreler:
    -------------
    kaynak_docx : Path
        Kaynak şablon dosyası
    esleme : dict[str, str]
        Placeholder eşleme dictionary'si
    gecici_dizin : Path
        Geçici dosyalar için dizin
    gunluk_ref : Logger, optional
        Logger referansı
    otomatik_takip : bool
        True ise geçici dosya otomatik temizlik için kaydedilir

    Döndürür:
    ---------
    Path
        Render edilmiş geçici dosya yolu
    """
    gecici_dizin.mkdir(parents=True, exist_ok=True)

    if not kaynak_docx.exists():
        hata = f"Şablon bulunamadı: {kaynak_docx}"
        if gunluk_ref:
            gunluk_ref.error(f"HATA: {hata}")
        raise FileNotFoundError(hata)

    cikti_yolu = gecici_dizin / f"{kaynak_docx.stem}__{uuid.uuid4().hex[:10]}.docx"

    if gunluk_ref:
        gunluk_ref.info(f"Render: {kaynak_docx.name} -> {cikti_yolu.name}")

    belge = DocxDocument(str(kaynak_docx))
    belgede_placeholder_degistir(belge, esleme, gunluk_ref)
    belge.save(str(cikti_yolu))

    # Otomatik temizlik için kaydet
    if otomatik_takip:
        _gecici_dosya_kaydet(cikti_yolu)

    return cikti_yolu


def belgeleri_birlestir(
    temel_docx: Path,
    eklenecekler: list[Path],
    cikti_docx: Path,
    gunluk_ref: Optional[logging.Logger] = None
) -> None:
    """
    Birden fazla belgeyi tek belgede birleştirir (Lego tarzı).

    Parametreler:
    -------------
    temel_docx : Path
        Temel belge (ilk belge)
    eklenecekler : list[Path]
        Eklenecek belgeler listesi
    cikti_docx : Path
        Çıktı belge yolu
    gunluk_ref : Logger, optional
        Logger referansı
    """
    if not temel_docx.exists():
        hata = f"Temel belge bulunamadı: {temel_docx}"
        if gunluk_ref:
            gunluk_ref.error(f"HATA: {hata}")
        raise FileNotFoundError(hata)

    if gunluk_ref:
        gunluk_ref.info(f"Belge birleştirme başladı (temel: {temel_docx.name})")

    temel = DocxDocument(str(temel_docx))
    birlestiricisi = Composer(temel)

    for i, yol in enumerate(eklenecekler, 1):
        yol = Path(yol)
        if not yol.exists():
            hata = f"Eklenecek belge bulunamadı: {yol}"
            if gunluk_ref:
                gunluk_ref.error(f"HATA: {hata}")
            raise FileNotFoundError(hata)

        if gunluk_ref:
            gunluk_ref.info(f"  [{i}/{len(eklenecekler)}] Ekleniyor: {yol.name}")

        birlestiricisi.append(DocxDocument(str(yol)))

    cikti_docx.parent.mkdir(parents=True, exist_ok=True)
    birlestiricisi.save(str(cikti_docx))

    if gunluk_ref:
        gunluk_ref.info(f"Birleştirilmiş belge kaydedildi: {cikti_docx.name}")


def global_placeholder_uygula(
    giris_docx: Path,
    global_esleme: dict[str, str],
    cikti_docx: Path,
    gunluk_ref: Optional[logging.Logger] = None
) -> None:
    """
    Mevcut belgeye global placeholder'ları uygular.

    Parametreler:
    -------------
    giris_docx : Path
        Giriş belge yolu
    global_esleme : dict[str, str]
        Global placeholder eşleme
    cikti_docx : Path
        Çıktı belge yolu
    gunluk_ref : Logger, optional
        Logger referansı
    """
    if gunluk_ref:
        gunluk_ref.info(f"Global placeholder'lar uygulanıyor: {giris_docx.name}")

    belge = DocxDocument(str(giris_docx))
    belgede_placeholder_degistir(belge, global_esleme, gunluk_ref)
    cikti_docx.parent.mkdir(parents=True, exist_ok=True)
    belge.save(str(cikti_docx))

    if gunluk_ref:
        gunluk_ref.info(f"Global placeholder'lar uygulandı: {cikti_docx.name}")


def footer_seri_numarasi_yaz(
    belge_yolu: Path,
    seri_numarasi: str,
    gunluk_ref: Optional[logging.Logger] = None
) -> None:
    """
    Belgenin footer'ına seri numarası yazar.

    {/SERIAL/} etiketini seri numarası ile değiştirir.

    Parametreler:
    -------------
    belge_yolu : Path
        İşlenecek belge yolu (yerinde değiştirilir)
    seri_numarasi : str
        Yazılacak seri numarası
    gunluk_ref : Logger, optional
        Logger referansı
    """
    try:
        belge = DocxDocument(str(belge_yolu))

        esleme = {"{/SERIAL/}": seri_numarasi}
        degistirilen = 0

        # Footer'larda ara
        for bolum in belge.sections:
            footer_listesi = [
                bolum.footer,
                bolum.first_page_footer,
                bolum.even_page_footer
            ]

            for footer in footer_listesi:
                for paragraf in _tum_paragraflari_yinele(footer):
                    onceki = "".join(r.text or "" for r in paragraf.runs)
                    if "{/SERIAL/}" in onceki:
                        _paragrafta_degistir(paragraf, esleme)
                        degistirilen += 1

        # Değişiklik yapıldıysa kaydet
        if degistirilen > 0:
            belge.save(str(belge_yolu))
            if gunluk_ref:
                gunluk_ref.info(f"Seri numarası footer'a yazıldı: {seri_numarasi} ({degistirilen} konum)")
        elif gunluk_ref:
            gunluk_ref.warning("Footer'da {/SERIAL/} etiketi bulunamadı")

    except Exception as e:
        if gunluk_ref:
            gunluk_ref.error(f"Footer'a seri numarası yazılamadı: {e}")
        raise


# =============================================================================
# Geriye Uyumluluk İçin Alias'lar
# =============================================================================

replace_placeholders_in_document = belgede_placeholder_degistir
render_to_temp = gecici_dosyaya_render_et
compose_docs = belgeleri_birlestir
apply_globals_to_docx = global_placeholder_uygula
