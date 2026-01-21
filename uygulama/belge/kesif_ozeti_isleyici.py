"""
Keşif Özeti İşleyici Modülü
===========================

Keşif özeti belgesi üzerinde opsiyon işaretleme ve numaralandırma işlemleri.

Bu modül, Word belgesindeki ürün bloklarını tanımlayarak:
- Opsiyonlu alt ürünleri işaretler (renklendirme + "OPSİYON" etiketi)
- Sıra numarası (NO) placeholder'larını doldurur
- Alt ürün bloklarını doğru konuma taşır
"""

from __future__ import annotations

import re
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Modül günlükleyicisi
gunluk = logging.getLogger(__name__)


# ============================================================
# REGEX PATTERN'LERİ
# ============================================================

# Ürün başlığı: {/<BASLIK>LB/}
BASLIK_TAG_RE = re.compile(
    r"\{/\s*<\s*BASLIK\s*>\s*([A-Za-z0-9_]+)\s*/\}",
    re.IGNORECASE
)

# Alt ürün başlığı: {/<ALTRN>LB_1/}
ALTRN_TAG_RE = re.compile(
    r"\{/\s*<\s*ALTRN\s*>\s*([A-Za-z0-9_]+)\s*_\s*([0-9]+)\s*/\}",
    re.IGNORECASE
)

# NO placeholder: {/NO/}
NO_TAG_RE = re.compile(r"\{/\s*NO\s*/\}", re.IGNORECASE)


# ============================================================
# YARDIMCI FONKSİYONLAR
# ============================================================

def row_text(row) -> str:
    """Satırdaki tüm hücre metinlerini birleştirir (okuma amaçlı)."""
    return "\n".join((c.text or "") for c in row.cells)


def iter_paragraphs_in_cell(cell):
    """Hücre içindeki tüm paragrafları ve iç tabloları dolaşır."""
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        for r in t.rows:
            for c in r.cells:
                yield from iter_paragraphs_in_cell(c)


def replace_in_paragraph_runs(paragraph, pattern: re.Pattern, repl: str) -> bool:
    """
    Format bozmadan replace (run-safe):
    - Run metinlerini birleştirir
    - Regex replace uygular
    - Sonucu ilk run'a yazar, diğer run'ları boşaltır
    
    Returns:
        bool: Değişiklik yapıldıysa True
    """
    runs = paragraph.runs
    if not runs:
        return False

    full = "".join(r.text or "" for r in runs)
    if not full:
        return False

    new = pattern.sub(repl, full)
    if new == full:
        return False

    runs[0].text = new
    for r in runs[1:]:
        r.text = ""
    return True


def paragraph_has_no(paragraph) -> bool:
    """Paragrafta {/NO/} placeholder'ı var mı?"""
    full = "".join(r.text or "" for r in paragraph.runs)
    return bool(NO_TAG_RE.search(full))


def set_cell_shading(cell, fill_hex: str) -> None:
    """Hücre arkaplan rengini ayarlar (hex, # olmadan)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())


def set_row_shading(row, fill_hex: str) -> None:
    """Satırdaki tüm hücrelere arkaplan rengi uygular."""
    for cell in row.cells:
        set_cell_shading(cell, fill_hex)


def shade_range(table, start: int, end: int, fill_hex: str) -> None:
    """Satır aralığına arkaplan rengi uygular."""
    for i in range(start, end + 1):
        set_row_shading(table.rows[i], fill_hex)


def replace_no_in_range(table, start: int, end: int, repl: str) -> bool:
    """
    Belirtilen satır aralığında {/NO/} placeholder'larını değiştirir.
    Format koruyarak replace yapar.
    
    Returns:
        bool: Herhangi bir değişiklik yapıldıysa True
    """
    changed_any = False
    for i in range(start, end + 1):
        row = table.rows[i]
        for cell in row.cells:
            for p in iter_paragraphs_in_cell(cell):
                if paragraph_has_no(p):
                    if replace_in_paragraph_runs(p, NO_TAG_RE, repl):
                        changed_any = True
    return changed_any


# ============================================================
# BLOK BULMA FONKSİYONLARI
# ============================================================

def find_product_blocks(table):
    """
    Ürün bloklarını {/<BASLIK>.../} etiketlerinden çıkarır.
    
    Returns:
        list[tuple]: [(PCODE, start_idx, end_idx), ...]
    """
    headers = []
    for i, row in enumerate(table.rows):
        m = BASLIK_TAG_RE.search(row_text(row))
        if m:
            code = m.group(1).strip()
            if code:
                headers.append((code.upper(), i))

    blocks = []
    for idx, (code, start) in enumerate(headers):
        next_start = headers[idx + 1][1] if idx + 1 < len(headers) else len(table.rows)
        end = next_start - 1
        blocks.append((code, start, end))
    return blocks


def find_product_header_trs(table):
    """
    Ürün başlık satırlarının _tr referansını döndürür.
    
    Returns:
        list[tuple]: [(PCODE, header_tr), ...]
    """
    headers = []
    for row in table.rows:
        m = BASLIK_TAG_RE.search(row_text(row))
        if m:
            code = m.group(1).strip().upper()
            headers.append((code, row._tr))
    return headers


def find_alt_blocks_in_range(table, product_code: str, start: int, end: int):
    """
    Ürün bloğu içinde {/<ALTRN>URUN_N/} satırlarını alt blok başlığı sayar.
    
    Returns:
        list[tuple]: [(ALT_KEY, astart, aend), ...] örn: "LB_1"
    """
    pc = product_code.upper()
    headers = []

    for i in range(start, end + 1):
        t = row_text(table.rows[i])
        m = ALTRN_TAG_RE.search(t)
        if not m:
            continue

        code = m.group(1).strip().upper()
        no = m.group(2).strip()

        if code == pc:
            headers.append((f"{code}_{no}", i))

    blocks = []
    for idx, (alt_key, astart) in enumerate(headers):
        next_start = headers[idx + 1][1] if idx + 1 < len(headers) else (end + 1)
        aend = next_start - 1
        blocks.append((alt_key, astart, aend))
    return blocks


def get_trs_for_range(table, start: int, end: int):
    """Satır aralığındaki w:tr objelerini sırayla döndürür."""
    return [table.rows[i]._tr for i in range(start, end + 1)]


# ============================================================
# SATIR TAŞIMA (ANCHOR-BASED)
# ============================================================

def move_trs_before_anchor(table, anchor_tr, trs_in_order: list) -> None:
    """
    trs_in_order satırlarını tablodan çıkarır ve anchor_tr'den HEMEN ÖNCE ekler.
    anchor_tr None ise tablo sonuna ekler.
    
    Not: anchor_tr bulunamazsa (çok nadir), sonuna ekler.
    """
    tbl = table._tbl

    # Hepsini önce çıkar
    for tr in trs_in_order:
        try:
            tbl.remove(tr)
        except ValueError:
            pass

    if anchor_tr is None:
        for tr in trs_in_order:
            tbl.append(tr)
        return

    # Anchor index bul
    try:
        anchor_idx = list(tbl).index(anchor_tr)
    except ValueError:
        # Anchor yoksa sonuna ekle
        for tr in trs_in_order:
            tbl.append(tr)
        return

    for off, tr in enumerate(trs_in_order):
        tbl.insert(anchor_idx + off, tr)


# ============================================================
# ANA LOJİK
# ============================================================

def opsiyon_ve_numaralama_uygula(
    belge_yolu: str | Path,
    cikti_yolu: str | Path,
    *,
    tablo_index: int = 0,
    opsiyon_alt_bayraklari: dict[str, bool] | None = None,
    opsiyon_urun_bayraklari: dict[str, bool] | None = None,
    opsiyon_renkleri: tuple[str, str] = ("FFF2CC", "D9E1F2"),
    logger: Optional[logging.Logger] = None
) -> Path:
    """
    Keşif özeti belgesi üzerinde opsiyon işaretleme ve numaralandırma işlemleri yapar.
    
    İşlem Adımları:
    1. Opsiyonlu alt blokları bulur
    2. {/NO/} -> "OPSİYON" değiştirir
    3. Arkaplan rengini değiştirir (dönüşümlü A/B renk)
    4. Alt blokları kendi ürününün sonuna taşır
    5. Tüm {/NO/} placeholder'larını sıralı numaralarla doldurur
    
    Parametreler:
    -------------
    belge_yolu : str | Path
        Kaynak Word belgesi yolu
    cikti_yolu : str | Path
        Çıktı belgesi yolu
    tablo_index : int
        İşlenecek tablonun indeksi (varsayılan: 0)
    opsiyon_alt_bayraklari : dict[str, bool]
        Alt ürün opsiyon bayrakları, örn: {"LB_1": True, "LB_2": False, ...}
    opsiyon_urun_bayraklari : dict[str, bool]
        Ürün bazlı opsiyon bayrakları, örn: {"LB": True, "ZR20": False}
    opsiyon_renkleri : tuple[str, str]
        Dönüşümlü renk çifti (hex, # olmadan), örn: ("FFF2CC", "D9E1F2")
    logger : Optional[logging.Logger]
        Günlük kaydedici (None ise modül günlükleyicisi kullanılır)
    
    Returns:
        Path: Oluşturulan belgenin yolu
    """
    log = logger or gunluk
    
    opsiyon_alt_bayraklari = {
        k.upper(): bool(v) for k, v in (opsiyon_alt_bayraklari or {}).items()
    }
    opsiyon_urun_bayraklari = {
        k.upper(): bool(v) for k, v in (opsiyon_urun_bayraklari or {}).items()
    }
    
    log.info("=" * 60)
    log.info("OPSİYON VE NUMARALAMA UYGULAMASI")
    log.info("=" * 60)
    log.info(f"Kaynak belge: {belge_yolu}")
    log.info(f"Çıktı belgesi: {cikti_yolu}")
    log.info(f"Tablo index: {tablo_index}")
    log.info(f"Opsiyon alt sayısı: {len(opsiyon_alt_bayraklari)}")
    log.info(f"Opsiyon ürün sayısı: {len(opsiyon_urun_bayraklari)}")

    doc = Document(str(belge_yolu))
    
    if tablo_index >= len(doc.tables):
        hata = f"Tablo index {tablo_index} belge içinde bulunamadı (toplam: {len(doc.tables)})"
        log.error(f"HATA: {hata}")
        raise ValueError(hata)
    
    table = doc.tables[tablo_index]
    log.info(f"Tablo yüklendi: {len(table.rows)} satır x {len(table.columns)} sütun")

    # Ürün blokları ve header sırası
    product_blocks = find_product_blocks(table)
    header_trs = find_product_header_trs(table)
    block_map = {code: (start, end) for code, start, end in product_blocks}
    
    log.info(f"Toplam {len(product_blocks)} ürün bloğu bulundu")
    for code, start, end in product_blocks:
        log.info(f"  - {code}: satır {start}-{end}")

    # 1) Opsiyon alt bloklarını işle
    log.info("=" * 60)
    log.info("ADIM 1: OPSİYON ALT BLOKLARI İŞLENİYOR")
    log.info("=" * 60)
    
    toplam_opsiyon = 0
    
    for idx in range(len(header_trs) - 1, -1, -1):
        pcode, _this_header_tr = header_trs[idx]
        pcode = pcode.upper()

        if pcode not in block_map:
            continue

        pstart, pend = block_map[pcode]

        # Anchor: bir sonraki ürün başlığı satırının tr'si
        anchor_tr = header_trs[idx + 1][1] if idx + 1 < len(header_trs) else None

        alt_blocks = find_alt_blocks_in_range(table, pcode, pstart, pend)
        product_is_opt = opsiyon_urun_bayraklari.get(pcode, False)

        opt_trs = []
        opt_block_color_counter = 0

        for alt_key, astart, aend in alt_blocks:
            is_opt = opsiyon_alt_bayraklari.get(alt_key.upper(), False) or product_is_opt
            if not is_opt:
                continue

            log.info(f"  Opsiyon işaretleniyor: {alt_key} (satır {astart}-{aend})")
            toplam_opsiyon += 1

            # NO -> OPSİYON (format koruyarak)
            replace_no_in_range(table, astart, aend, "OPSİYON")

            # Arkaplan boya (blok bazında dönüşümlü)
            color = opsiyon_renkleri[opt_block_color_counter % 2]
            opt_block_color_counter += 1
            shade_range(table, astart, aend, color)

            # Bu alt bloğun satırlarını sırayla topla
            opt_trs.extend(get_trs_for_range(table, astart, aend))

        if opt_trs:
            log.info(f"  {len(opt_trs)} satır taşınıyor (ürün: {pcode})")
            # KRİTİK: sonraki ürün başlığından HEMEN ÖNCE
            move_trs_before_anchor(table, anchor_tr, opt_trs)

            # Taşıma sonrası indexler değişir: block_map'i tazele
            product_blocks = find_product_blocks(table)
            block_map = {code: (start, end) for code, start, end in product_blocks}

    log.info(f"✓ Toplam {toplam_opsiyon} opsiyon bloğu işlendi")

    # 2) Global NO numaralandırma
    log.info("=" * 60)
    log.info("ADIM 2: SIRA NUMARALARI DOLDURULUYOR")
    log.info("=" * 60)
    
    no_counter = 1
    opsiyon_sayisi = 0
    
    for i, row in enumerate(table.rows):
        t = row_text(row).upper()

        # Opsiyon satırları: NO kaldıysa OPSİYON yap, numara verme
        if "OPSİYON" in t or "OPSIYON" in t:
            for cell in row.cells:
                for p in iter_paragraphs_in_cell(cell):
                    replace_in_paragraph_runs(p, NO_TAG_RE, "OPSİYON")
            opsiyon_sayisi += 1
            continue

        # Normal satır: NO placeholder varsa sırayı yaz
        replaced_any = False
        for cell in row.cells:
            for p in iter_paragraphs_in_cell(cell):
                if paragraph_has_no(p):
                    if replace_in_paragraph_runs(p, NO_TAG_RE, str(no_counter)):
                        replaced_any = True

        if replaced_any:
            no_counter += 1

    log.info(f"✓ {no_counter - 1} sıra numarası yazıldı")
    log.info(f"✓ {opsiyon_sayisi} opsiyon satırı işaretlendi")

    # Belgeyi kaydet
    cikti_yolu = Path(cikti_yolu)
    cikti_yolu.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(cikti_yolu))
    
    log.info("=" * 60)
    log.info(f"✓ Belge başarıyla kaydedildi: {cikti_yolu}")
    log.info("=" * 60)
    
    return cikti_yolu
