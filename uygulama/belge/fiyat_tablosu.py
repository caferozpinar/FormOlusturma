# uygulama/belge/fiyat_tablosu.py
"""
Fiyat Tablosu Oluşturma Modülü
===============================

Word belgesinde fiyat tablosu oluşturma ve doldurma işlemleri.

Bu modül artık standart Türk para formatını (xxx.xxx.xxx,xx) kullanır.
Tüm fiyat işlemleri FiyatFormatlayici üzerinden yapılır.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from uygulama.yardimcilar.fiyat_formatlayici import FiyatFormatlayici


# Opsiyon renkleri (varsayılan değerler)
_OPSIYON_RENKLERI = [
    (173, 216, 230),  # Açık Mavi
    (255, 218, 185),  # Açık Turuncu
]


def _renk_config_oku(config_yolu: Optional[Path] = None) -> list[tuple[int, int, int]]:
    """
    Opsiyon renk ayarlarını config.txt dosyasından okur.
    
    Config formatı:
    {{OPSIYON_RENK_1}}="173,216,230"
    {{OPSIYON_RENK_2}}="255,218,185"
    
    Parametreler:
    -------------
    config_yolu : Optional[Path]
        Config dosya yolu (None ise varsayılan config.txt kullanılır)
    
    Döndürür:
    ---------
    list[tuple[int, int, int]]
        RGB renk listesi
    """
    if config_yolu is None:
        # Varsayılan: ./config.txt
        config_yolu = Path("config.txt")
    
    if not config_yolu.exists():
        return _OPSIYON_RENKLERI
    
    renkler = []
    try:
        with open(config_yolu, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                
                # Format: {{OPSIYON_RENK_1}}="173,216,230"
                if '{{OPSIYON_RENK_' in line and '}}' in line:
                    # Değeri çıkar: "173,216,230"
                    if '=' in line:
                        value_part = line.split('=', 1)[1].strip()
                        # Tırnakları temizle
                        value_part = value_part.strip('"').strip("'")
                        
                        # RGB parse et
                        parts = value_part.split(',')
                        if len(parts) == 3:
                            r, g, b = int(parts[0]), int(parts[1]), int(parts[2])
                            renkler.append((r, g, b))
        
        return renkler if renkler else _OPSIYON_RENKLERI
    
    except Exception:
        return _OPSIYON_RENKLERI


def _satir_arka_plan_renklendir(row, rgb: tuple[int, int, int]) -> None:
    """
    Tablo satırının tüm hücrelerine arka plan rengi uygular.
    
    Parametreler:
    -------------
    row : docx.table.Row
        Tablo satırı
    rgb : tuple[int, int, int]
        RGB renk değeri (0-255)
    """
    for cell in row.cells:
        # Hücre shading öğesi al veya oluştur
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # Shading öğesi oluştur
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
        shd.set(qn('w:val'), 'clear')
        
        # Eski shading varsa kaldır
        for old_shd in tcPr.findall(qn('w:shd')):
            tcPr.remove(old_shd)
        
        # Yeni shading ekle
        tcPr.append(shd)


def _replace_in_table_preserve_runs(table, repl: dict[str, str]) -> None:
    """Tablodaki tüm hücrelerde placeholder değiştirme yapar."""
    for row in table.rows:
        for cell in row.cells:
            _replace_in_cell_preserve_runs(cell, repl)


def _urun_satir_indeksleri_ve_opsiyonlar(durum: dict[str, Any]) -> tuple[list[int], set[int]]:
    """
    adet_line_* veya brmfiyat_line_* alanlarında dolu olan indexleri döndürür.
    Ayrıca hangi satırların opsiyon olduğunu da belirler.
    
    Parametreler:
    -------------
    durum : dict[str, Any]
        Ürün formu verileri
    
    Döndürür:
    ---------
    tuple[list[int], set[int]]
        (Dolu satır indeksleri, Opsiyon satır indeksleri)
    """
    pat_adet = re.compile(r"^adet_line_(\d+)$")
    pat_fiyat = re.compile(r"^brmfiyat_line_(\d+)$")
    pat_ops = re.compile(r"^urun_ops_(\d+)$")
    
    idx: set[int] = set()
    opsiyon_idx: set[int] = set()

    for k, v in (durum or {}).items():
        # adet_line_* kontrol et
        m = pat_adet.match(str(k))
        if m and str(v).strip():
            idx.add(int(m.group(1)))
            continue

        # brmfiyat_line_* kontrol et
        m = pat_fiyat.match(str(k))
        if m and str(v).strip():
            idx.add(int(m.group(1)))
            continue
        
        # urun_ops_* checkbox kontrol et
        m = pat_ops.match(str(k))
        if m:
            # Checkbox işaretli mi?
            if v is True or str(v).lower() in ('true', '1', 'yes', 'checked'):
                opsiyon_idx.add(int(m.group(1)))

    return sorted(idx), opsiyon_idx


def _replace_in_cell_preserve_runs(cell, repl: dict[str, str]) -> None:
    """Hücredeki run'ları koruyarak placeholder değiştirme yapar."""
    for p in cell.paragraphs:
        if not p.runs:
            continue
        full = "".join(r.text for r in p.runs)
        new_full = full
        for k, v in repl.items():
            if k in new_full:
                new_full = new_full.replace(k, v)
        if new_full != full:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ""


def _replace_in_row_preserve_runs(row, repl: dict[str, str]) -> None:
    """Satırdaki tüm hücrelerde placeholder değiştirme yapar."""
    for cell in row.cells:
        _replace_in_cell_preserve_runs(cell, repl)


def _row_index_by_tr(table, tr) -> int:
    """Verilen TR elementinin tablodaki satır indeksini bulur."""
    trs = [r._tr for r in table.rows]
    return trs.index(tr)


def _insert_clone_before_tr(table, anchor_tr, clone_row_idx: int):
    """Belirtilen satırı klonlayarak anchor'dan önce ekler."""
    tbl = table._tbl
    anchor_idx = _row_index_by_tr(table, anchor_tr)
    tr_clone = deepcopy(table.rows[clone_row_idx]._tr)
    tbl.insert(anchor_idx, tr_clone)
    return table.rows[anchor_idx]


def _rows_and_totals_from_ui(
    urun_kodlari: list[str],
    oturum_onbellegi: dict[str, dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, str]]:
    """
    UI sırasına göre satırları flatten eder ve toplamları hesaplar.
    
    ÖNEMLİ: 
    - Opsiyon ürünler normal ürünlerden SONRA gelir
    - Opsiyon ürünler genel toplama DAHİL EDİLMEZ
    - Opsiyon ürünlerin NO alanı "OPSİYON" olarak ayarlanır

    Parametreler:
    -------------
    urun_kodlari : list[str]
        Ürün kodları listesi
    oturum_onbellegi : dict[str, dict[str, Any]]
        Oturum verileri

    Döndürür:
    ---------
    tuple[list[dict], list[dict], dict]
        (Normal ürün satırları, Opsiyon ürün satırları, Toplam verileri)
    """
    normal_rows: list[dict[str, Any]] = []
    opsiyon_rows: list[dict[str, Any]] = []

    toplam_f = 0.0
    kdv_oran_text: Optional[str] = None

    for urun in urun_kodlari:
        durum = oturum_onbellegi.get(urun, {})
        idxler, opsiyon_idxler = _urun_satir_indeksleri_ve_opsiyonlar(durum)

        # KDV oranını ilk bulduğumuz dolu alandan çek
        if kdv_oran_text is None:
            t = str(durum.get("kdv_line", "")).strip()
            if t:
                kdv_oran_text = t

        for i in idxler:
            # Opsiyon mu?
            is_opsiyon = i in opsiyon_idxler

            # Adet bilgisi
            adet = str(durum.get(f"adet_line_{i}", "")).strip()

            # Ürün label
            urun_adi = str(durum.get(f"urun_label_{i}", "")).strip()
            if not urun_adi:
                urun_adi = urun

            # NO değeri
            if is_opsiyon:
                no_degeri = "OPSİYON"
            else:
                no_degeri = str(durum.get(f"no_line_{i}", "")).strip()
                if not no_degeri:
                    no_degeri = str(i)

            # Fiyat bilgileri
            br_str = str(durum.get(f"brmfiyat_line_{i}", "")).strip()
            tp_str = str(durum.get(f"top_line_{i}", "")).strip()

            # Float'a çevir
            br_float = FiyatFormatlayici.turk_format_to_float(br_str)
            tp_float = FiyatFormatlayici.turk_format_to_float(tp_str)

            # Sadece NORMAL ürünleri toplama ekle
            if not is_opsiyon:
                toplam_f += tp_float

            # Türk formatına çevir
            br_formatli = FiyatFormatlayici.float_to_turk_format(br_float)
            tp_formatli = FiyatFormatlayici.float_to_turk_format(tp_float)

            row_data = {
                "NO": no_degeri,
                "URUNAD": adet,
                "URUN": urun_adi,
                "BRFY": br_formatli,
                "TPFY": tp_formatli,
                "is_opsiyon": is_opsiyon,  # Renklendirme için flag
            }

            # Normal veya opsiyon listesine ekle
            if is_opsiyon:
                opsiyon_rows.append(row_data)
            else:
                normal_rows.append(row_data)

    # KDV hesaplama (sadece normal ürünler için)
    kdv_f = 0.0
    kdv_oran_formatli = ""

    if kdv_oran_text:
        kdv_oran_temiz = kdv_oran_text.replace("%", "").strip()
        oran_float = FiyatFormatlayici.turk_format_to_float(kdv_oran_temiz)
        kdv_f = toplam_f * (oran_float / 100.0)
        kdv_oran_formatli = FiyatFormatlayici.float_to_turk_format(oran_float)

    # Toplamları formatla
    totals = {
        "TPFY": FiyatFormatlayici.float_to_turk_format(toplam_f),
        "KDVFY": FiyatFormatlayici.float_to_turk_format(kdv_f),
        "TOPKDV": FiyatFormatlayici.float_to_turk_format(toplam_f + kdv_f),
        "KDV": kdv_oran_formatli,
    }

    return normal_rows, opsiyon_rows, totals


def fiyat_tablosu_uret_ve_doldur(
    fiyat_sablon_yolu: Path,
    cikti_yolu: Path,
    urun_kodlari: list[str],
    oturum_onbellegi: dict[str, dict[str, Any]],
    *,
    baslik_metni: str,
    table_idx: int,
    template_row_idx: int,
    total_anchor_row_idx: int,
    logger=None,
    renk_config_yolu: Optional[Path] = None,
    # Şablondaki placeholder isimlerini buradan tek noktadan değiştir:
    ph_no: str = "{/NO/}",
    ph_urunad: str = "{/URUNAD/}",  # Adet
    ph_urun: str = "{/URUN/}",      # Ürün label
    ph_brfy: str = "{/URUNBRFY/}",  # Birim fiyat
    ph_tpfy: str = "{/URUNTPFY/}",  # Toplam fiyat
    ph_toplam: str = "{/TPFY/}",
    ph_kdvfy: str = "{/KDVFY/}",
    ph_toplamkdv: str = "{/TOPKDV/}",
    ph_kdv_oran: str = "{/KDV/}",
    ph_urunler: str = "{/URUNLER/}",  # Başlık metni için
) -> None:
    """
    Fiyat tablosunu oluşturur ve doldurur.

    YENİ ÖZELLİKLER:
    - Opsiyon ürünler normal ürünlerden SONRA listelenir
    - Opsiyon ürünler genel toplama DAHİL EDİLMEZ
    - Opsiyon ürünlerin arka planı renklendirilir (mavi-turuncu örgüsü)
    - Opsiyon ürünlerin NO alanı "OPSİYON" yazılır

    Parametreler:
    -------------
    fiyat_sablon_yolu : Path
        Fiyat tablosu şablon dosyası
    cikti_yolu : Path
        Çıktı dosyası
    urun_kodlari : list[str]
        Ürün kodları listesi (örn: ['LK', 'ZTK'])
    oturum_onbellegi : dict[str, dict[str, Any]]
        Her ürün için oturum verileri
        Beklenen alanlar:
        - adet_line_1..8: Ürün adedi
        - urun_label_1..8: Ürün açıklaması
        - brmfiyat_line_1..8: Birim fiyat (Türk formatı)
        - top_line_1..8: Toplam fiyat (Türk formatı)
        - urun_ops_1..8: Opsiyon checkbox (boolean)
        - kdv_line: KDV oranı (Türk formatı)
    baslik_metni : str
        Tablo başlığı metni ({/URUNLER/} placeholder'ına yazılır)
    table_idx : int
        Tablo indeksi (şablonda birden fazla tablo varsa)
    template_row_idx : int
        Şablon satır indeksi (kopyalanacak satır)
    total_anchor_row_idx : int
        Toplam satırı indeksi (anchor)
    logger : Optional[Logger]
        Logger nesnesi
    renk_config_yolu : Optional[Path]
        Renk config dosya yolu (None ise varsayılan)
    ph_* : str
        Placeholder isimleri (özelleştirilebilir)
    """

    # Satırları ve toplamları al (normal + opsiyon ayrı)
    normal_rows, opsiyon_rows, totals = _rows_and_totals_from_ui(urun_kodlari, oturum_onbellegi)
    
    # Tüm satırları birleştir: ÖNCE normal, SONRA opsiyon
    all_rows = normal_rows + opsiyon_rows

    if logger:
        logger.info(f"Fiyat tablosu: {len(normal_rows)} normal, {len(opsiyon_rows)} opsiyon satır")
        logger.debug(f"Toplam (sadece normal): {totals['TPFY']}, KDV: {totals['KDVFY']}, "
                    f"Genel Toplam: {totals['TOPKDV']}")

    doc = Document(str(fiyat_sablon_yolu))
    table = doc.tables[table_idx]

    # TOPLAM satırını anchor al
    anchor_tr = table.rows[total_anchor_row_idx]._tr

    # 1) Satır sayısını ayarla
    hedef = max(1, len(all_rows))
    mevcut = 1
    eklenecek = max(0, hedef - mevcut)

    for _ in range(eklenecek):
        _insert_clone_before_tr(table, anchor_tr, template_row_idx)

    # 2) Ürün satırlarını bul
    urun_satirlari = []
    for r in table.rows:
        if r._tr is anchor_tr:
            break
        urun_satirlari.append(r)

    urun_satirlari = urun_satirlari[template_row_idx : template_row_idx + len(all_rows)]

    # 3) Renk konfigürasyonunu oku
    renkler = _renk_config_oku(renk_config_yolu)
    opsiyon_renk_idx = 0

    # 4) Satır satır doldur VE renklendir
    for row_obj, data in zip(urun_satirlari, all_rows):
        # Placeholder'ları doldur
        repl = {
            ph_no: data.get("NO", ""),
            ph_urunad: data.get("URUNAD", ""),
            ph_urun: data.get("URUN", ""),
            ph_brfy: data.get("BRFY", ""),
            ph_tpfy: data.get("TPFY", ""),
        }
        _replace_in_row_preserve_runs(row_obj, repl)

        # Opsiyon ise arka planı renklendir
        if data.get("is_opsiyon", False):
            renk = renkler[opsiyon_renk_idx % len(renkler)]
            _satir_arka_plan_renklendir(row_obj, renk)
            opsiyon_renk_idx += 1
            
            if logger:
                logger.debug(f"Opsiyon satır renklendirildi: {data.get('URUN')} - RGB{renk}")

    # 5) TOPLAM satırını doldur
    anchor_idx = _row_index_by_tr(table, anchor_tr)
    anchor_row = table.rows[anchor_idx]

    _replace_in_row_preserve_runs(
        anchor_row,
        {
            ph_toplam: totals.get("TPFY", ""),
            ph_kdv_oran: totals.get("KDV", ""),
        },
    )

    # 6) KDV ve Genel Toplam satırları
    if anchor_idx + 1 < len(table.rows):
        _replace_in_row_preserve_runs(
            table.rows[anchor_idx + 1],
            {ph_kdvfy: totals.get("KDVFY", "")}
        )
    if anchor_idx + 2 < len(table.rows):
        _replace_in_row_preserve_runs(
            table.rows[anchor_idx + 2],
            {ph_toplamkdv: totals.get("TOPKDV", "")}
        )

    # 7) Yedek: Tüm tabloda kalan placeholder'ları doldur
    _replace_in_table_preserve_runs(
        table,
        {
            ph_toplam: totals.get("TPFY", ""),
            ph_kdvfy: totals.get("KDVFY", ""),
            ph_toplamkdv: totals.get("TOPKDV", ""),
            ph_kdv_oran: totals.get("KDV", ""),
            ph_urunler: baslik_metni,
        },
    )

    doc.save(str(cikti_yolu))

    if logger:
        logger.info(f"✓ Fiyat tablosu yazıldı: {cikti_yolu.name}")
        logger.info(f"  - Normal ürünler: {len(normal_rows)}")
        logger.info(f"  - Opsiyon ürünler: {len(opsiyon_rows)}")
        logger.info(f"  - Genel toplam (KDV dahil, sadece normal): {totals['TOPKDV']}")
